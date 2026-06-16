#!/usr/bin/env python3
"""Genera MANUAL_USUARIO_RELATIC_CERTIFICADOS.docx desde contenido operativo actual."""

from __future__ import annotations

import os
import sys

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', '..'))
DOC_NAME = 'MANUAL_USUARIO_RELATIC_CERTIFICADOS.docx'
OUT_FILE = os.path.join(ROOT, 'docs', DOC_NAME)
OUT_STATIC = os.path.join(ROOT, 'static', 'downloads', DOC_NAME)


def _h(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def _p(doc: Document, text: str, bold: bool = False) -> None:
    run = doc.add_paragraph().add_run(text)
    run.bold = bold


def _bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style='List Bullet')


def _table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
    doc.add_paragraph()


def build() -> str:
    os.makedirs(os.path.dirname(OUT_STATIC), exist_ok=True)
    doc = Document()
    title = doc.add_heading('Manual de usuario — Certificados Relatic', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = sub.add_run('Easy NodeOne · Relatic Panamá\nVersión 2.0 — Junio 2026 (Nivel 2A visual + 2B portal usuario)')
    r.font.size = Pt(11)

    _h(doc, '1. Alcance', 1)
    _p(doc, 'Este manual cubre certificados de EVENTOS (cursos, seminarios, congresos). No aplica a cursos académicos del módulo Educación.')
    _bullet(doc, 'Administradores: emitir, descargar, revocar y compartir certificados.')
    _bullet(doc, 'Usuarios finales: ver y descargar sus certificados en Mis Certificados.')

    _h(doc, '2. Accesos en la plataforma', 1)
    _table(
        doc,
        ['Rol', 'Menú / ruta', 'URL'],
        [
            ['Admin', 'Comercial → Eventos → Participantes / Certificados', '/admin/events/<id>/participants'],
            ['Admin', 'Operaciones → Certificados → Eventos / Plantillas', '/admin/certificate-events'],
            ['Usuario', 'Operaciones → Mis Certificados (debajo de Certificados)', '/certificates'],
            ['Usuario', 'Perfil → Mis Certificados', '/certificates'],
            ['Público', 'Verificación QR / código', '/certificates/verify/<código>'],
        ],
    )

    _h(doc, '3. Flujo administrador (emisión)', 1)
    steps = [
        'Crear o abrir el evento en Eventos.',
        'Cargar participantes (ver sección 4).',
        'Marcar asistencia: Check-in (o tipo Revisor).',
        'Generar certificado: botón certificado en la fila o pantalla Certificados del evento.',
        'Descargar PDF (admin) o dejar que el usuario lo descargue en Mis Certificados (2B).',
        'Opcional: revocar y volver a generar si hubo error.',
    ]
    for i, s in enumerate(steps, 1):
        _bullet(doc, f'{i}. {s}')

    _h(doc, '4. Campos del participante (pantalla y base de datos)', 1)
    _table(
        doc,
        ['Campo', 'Descripción', 'Valores / notas'],
        [
            ['Primer nombre', 'Obligatorio', 'Texto'],
            ['Segundo nombre', 'Opcional', 'Texto'],
            ['Primer apellido', 'Obligatorio', 'Texto'],
            ['Segundo apellido', 'Opcional', 'Texto'],
            ['Documento (cédula/pasaporte)', 'Recomendado', 'document_id'],
            ['Email', 'Recomendado para portal 2B', 'Debe coincidir con cuenta EN1 para descarga automática'],
            ['Teléfono', 'Opcional', 'Texto'],
            ['Tipo participante', 'participant_type', 'external, member, reviewer, invited, speaker, staff'],
            ['Estado de pago', 'payment_status', 'pending, paid, complimentary, waived, not_required'],
            ['Asistencia', 'attendance_status', 'pending, checked_in, attended, absent'],
            ['Estado certificado', 'certificate_status', 'pending, issued, generated, sent'],
            ['Origen', 'registration_source', 'admin_manual, import, inscripcion_confirmada'],
            ['Usuario EN1', 'user_id', 'Vinculado si email coincide al entrar a Mis Certificados'],
            ['Notas', 'notes', 'Texto libre'],
        ],
    )

    _h(doc, '5. Importación Excel — columnas A a J', 1)
    _p(doc, 'Plantilla «LISTA PARA CERTIFICADOS». Fila 1 puede ser encabezado (se detecta automáticamente).')
    _table(
        doc,
        ['Col', 'Campo', 'Obligatorio'],
        [
            ['A', 'Primer nombre', 'Sí'],
            ['B', 'Segundo nombre', 'No'],
            ['C', 'Primer apellido', 'Sí'],
            ['D', 'Segundo apellido', 'No'],
            ['E', 'Documento', 'No'],
            ['F', 'Email', 'No (recomendado)'],
            ['G', 'Teléfono', 'No'],
            ['H', 'Tipo participante', 'No — external, reviewer, member, etc.'],
            ['I', 'Estado de pago', 'No — pending, paid, complimentary, waived, not_required'],
            ['J', 'Notas', 'No'],
        ],
    )
    _p(doc, 'Menú Agregar → Importar Excel, o Agregar → Desde registros confirmados (inscripciones con pago confirmado).')

    _h(doc, '6. Pantalla Participantes — columnas visibles', 1)
    _table(
        doc,
        ['Columna', 'Contenido'],
        [
            ['Participante', 'Nombre, documento, tipo (Miembro/Revisor/Externo), badge EN1 si vinculado, origen'],
            ['Contacto', 'Email y teléfono'],
            ['Asistencia', 'Pendiente / Check-in / Ausente'],
            ['Certificado', 'Emitido + código, o Pendiente'],
            ['Acciones', 'Check-in, ausente, pendiente, editar, certificados, PDF, verificar, generar, eliminar'],
        ],
    )

    _h(doc, '7. Pantalla Certificados del evento (admin)', 1)
    _bullet(doc, 'Estadísticas: emitidos, pendientes elegibles, revocados, participantes.')
    _bullet(doc, 'Generar todos los elegibles o seleccionados con casilla.')
    _bullet(doc, 'Tabla emitidos: código, participante, fecha, estado, acciones PDF / verificar / revocar.')

    _h(doc, '8. Certificado emitido — campos', 1)
    _table(
        doc,
        ['Campo', 'Descripción'],
        [
            ['certificate_number', 'Código único (ej. EN1-2026-000001, REV-2026-… para revisores)'],
            ['certificate_type', 'participation o reviewer'],
            ['title', 'Título en PDF'],
            ['issued_date', 'Fecha de emisión'],
            ['status', 'generated, sent, revoked'],
            ['is_active', 'true/false — revocados no aparecen en portal usuario'],
            ['certificate_url', 'Ruta al PDF'],
            ['verification_token', 'Token interno de verificación'],
            ['qr_path', 'Imagen QR en PDF'],
        ],
    )

    _h(doc, '9. Portal usuario — Mis Certificados (Nivel 2B)', 1)
    _p(doc, 'Ruta: /certificates')
    _h(doc, '9.1 Certificados de eventos', 2)
    _bullet(doc, 'Tarjetas con: tipo, evento, fechas, código, fecha emisión, estado.')
    _bullet(doc, 'Descargar PDF — solo si el admin ya generó el certificado.')
    _bullet(doc, 'Verificar — abre página pública de autenticidad.')
    _bullet(doc, 'Requisito: el email del participante debe coincidir con la cuenta (o user_id vinculado).')

    _h(doc, '9.2 Certificados de membresía', 2)
    _bullet(doc, 'Tipos REG (registro) y MEM (membresía activa).')
    _bullet(doc, 'Solicitar si cumple requisitos; Descargar PDF si ya fue emitido.')
    _bullet(doc, 'Código formato MEM-… o REG-… según configuración del tenant.')

    _h(doc, '10. Elegibilidad para generar certificado', 1)
    _bullet(doc, 'Participante con check-in o attended, O tipo reviewer.')
    _bullet(doc, 'Sin certificado activo previo para el mismo participante y evento.')
    _bullet(doc, 'Certificados revocados no bloquean una nueva emisión tras revocar.')

    _h(doc, '11. PDF institucional (Nivel 2A)', 1)
    _bullet(doc, 'Formato horizontal Letter con logos, firmas, sello, QR y URL de verificación.')
    _bullet(doc, 'Configuración por defecto: relatic_event_certificate_layout.json.')
    _bullet(doc, 'Override por evento: JSON en campo Plantilla certificado del evento.')

    _h(doc, '12. Verificación pública', 1)
    _p(doc, 'URL: https://apps.relatic.org/certificates/verify/<código>')
    _bullet(doc, 'Muestra válido / revocado, nombre participante y evento.')
    _bullet(doc, 'El QR del PDF apunta a esta URL.')

    _h(doc, '13. Resumen de responsabilidades', 1)
    _table(
        doc,
        ['Tarea', 'Admin', 'Usuario'],
        [
            ['Cargar participantes', 'Sí', 'No'],
            ['Marcar asistencia', 'Sí', 'No'],
            ['Generar certificado', 'Sí', 'No'],
            ['Descargar PDF evento', 'Sí', 'Sí (Mis Certificados)'],
            ['Solicitar cert. membresía', 'No', 'Sí (si aplica)'],
            ['Revocar certificado', 'Sí', 'No'],
            ['Verificar autenticidad', 'Ambos', 'Ambos'],
        ],
    )

    _h(doc, '14. Soporte y actualizaciones', 1)
    _p(doc, 'Documento generado desde dev/app. Descarga: /static/downloads/MANUAL_USUARIO_RELATIC_CERTIFICADOS.docx o docs/MANUAL_USUARIO_RELATIC_CERTIFICADOS.docx en el repositorio.')
    _p(doc, 'Referencia técnica: docs/MANUAL_OPERATIVO_RELATIC_CERTIFICADOS_EVENTOS.md y docs/RELATIC_CERTIFICADOS_PORTAL_USUARIO_2B.md')

    doc.save(OUT_FILE)
    import shutil

    shutil.copy2(OUT_FILE, OUT_STATIC)
    return OUT_FILE


if __name__ == '__main__':
    path = build()
    print(path)
    sys.exit(0)
